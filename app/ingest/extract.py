"""Tekstextractie uit brondocumenten (PDF, Word, Excel, platte tekst).

Elke extractor geeft platte tekst terug (inclusief tabellen, waar mogelijk
als eenvoudige rij-per-rij weergave). Bij fouten wordt een `ExtractionError`
gegooid met een gebruiksvriendelijke, Nederlandstalige melding.
"""
from __future__ import annotations

from pathlib import Path


class ExtractionError(Exception):
    """Duidelijke, aan de gebruiker te tonen foutmelding bij extractie."""


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".xlsx", ".xls"):
        return _extract_excel(path)
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ExtractionError(f"Bestandstype '{ext}' wordt niet ondersteund.")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as e:
        raise ExtractionError("pdfplumber is niet geïnstalleerd (zie requirements.txt).") from e

    parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                parts.append(f"--- Pagina {i} ---\n{text}")
                for table in page.extract_tables() or []:
                    rows = ["\t".join(c or "" for c in row) for row in table]
                    if rows:
                        parts.append("[Tabel]\n" + "\n".join(rows))
    except Exception as e:  # noqa: BLE001 - alle pdfplumber-fouten worden hier afgevangen
        raise ExtractionError(
            f"Kon '{path.name}' niet lezen als PDF ({e}). Mogelijk is dit een gescande "
            "(afbeelding-)PDF zonder doorzoekbare tekst — probeer eerst OCR toe te passen."
        ) from e

    full_text = "\n\n".join(parts).strip()
    if not full_text or len(full_text) < 20:
        raise ExtractionError(
            f"'{path.name}' bevat geen doorzoekbare tekst (waarschijnlijk een gescande PDF). "
            "Probeer het bestand eerst door OCR-software te halen en upload opnieuw."
        )
    return full_text


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as e:
        raise ExtractionError("python-docx is niet geïnstalleerd (zie requirements.txt).") from e

    try:
        document = docx.Document(str(path))
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"Kon '{path.name}' niet openen als Word-document ({e}).") from e

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        rows = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
        if rows:
            parts.append("[Tabel]\n" + "\n".join(rows))

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ExtractionError(f"'{path.name}' lijkt leeg te zijn (geen tekst gevonden).")
    return full_text


def _extract_excel(path: Path) -> str:
    try:
        import openpyxl
    except ImportError as e:
        raise ExtractionError("openpyxl is niet geïnstalleerd (zie requirements.txt).") from e

    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"Kon '{path.name}' niet openen als Excel-bestand ({e}).") from e

    parts: list[str] = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(v is not None for v in row):
                rows.append("\t".join("" if v is None else str(v) for v in row))
        if rows:
            parts.append(f"[Werkblad: {sheet.title}]\n" + "\n".join(rows))

    full_text = "\n\n".join(parts).strip()
    if not full_text:
        raise ExtractionError(f"'{path.name}' bevat geen gegevens (leeg werkboek).")
    return full_text


# ---------------------------------------------------------- typelabel-suggestie

_KEYWORD_TYPES = [
    ("offerte-aanvraag", ["offerteaanvraag", "aanbestedingsleidraad", "rfq", "uitvraag", "vraagspecificatie"]),
    ("SLA", ["service level agreement", "sla ", "sla-", "dienstverleningsovereenkomst"]),
    ("contractvoorwaarden", ["algemene voorwaarden", "contractvoorwaarden", "overeenkomst"]),
    ("PvE", ["programma van eisen", "pve", "eisenpakket"]),
]


def suggest_document_type(filename: str, text: str = "") -> str:
    haystack = f"{filename.lower()} {text[:3000].lower()}"
    for label, keywords in _KEYWORD_TYPES:
        if any(k in haystack for k in keywords):
            return label
    return "overig"
