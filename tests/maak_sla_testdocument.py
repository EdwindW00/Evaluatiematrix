"""Genereert een synthetisch SLA-testdocument met specifieke, alleen-hier-genoemde eisen,
om te verifiëren dat AI-matrixgeneratie informatie uit de SLA correct meeneemt."""
from pathlib import Path

import docx

OUT = Path(__file__).parent / "sample_sla.docx"

d = docx.Document()
d.add_heading("Service Level Agreement — Schoonmaakdienstverlening Kantoorlocatie Amsterdam", level=1)
d.add_paragraph(
    "Deze SLA maakt onderdeel uit van de overeenkomst tussen Opdrachtgever en Opdrachtnemer "
    "en specificeert de dienstverleningsniveaus, responstijden en boeteclausules."
)
d.add_heading("Artikel 3 — Responstijden en beschikbaarheid", level=2)
d.add_paragraph(
    "3.1 Opdrachtnemer garandeert een reactietijd van maximaal 2 uur op meldingen van "
    "calamiteiten (bijv. wateroverlast, ernstige vervuiling) gedurende kantooruren (08:00-18:00)."
)
d.add_paragraph(
    "3.2 Voor reguliere klachten geldt een maximale afhandeltermijn van 24 uur."
)
d.add_paragraph(
    "3.3 Opdrachtnemer dient te beschikken over een 24/7 bereikbare storingsdienst, "
    "aantoonbaar middels een telefoonnummer en escalatieprotocol."
)
d.add_heading("Artikel 5 — Kwaliteitscontrole en KPI's", level=2)
d.add_paragraph(
    "5.1 Opdrachtnemer voert minimaal maandelijks een kwaliteitsaudit uit conform de "
    "VSR-kwaliteitsmeetmethode (Vereniging Schoonmaak Research), met een minimale score van 7,5."
)
d.add_paragraph(
    "5.2 Bij een auditscore lager dan 7,0 gedurende twee opeenvolgende metingen is Opdrachtgever "
    "gerechtigd een boete van 5% van de maandfactuur in te houden."
)
d.add_heading("Artikel 7 — Personeel en continuïteit", level=2)
d.add_paragraph(
    "7.1 Opdrachtnemer garandeert een vast team van minimaal 3 medewerkers op de locatie, "
    "met een maximale personeelswisseling van 20% per jaar."
)
d.add_paragraph(
    "7.2 Alle medewerkers dienen te beschikken over een geldige VOG (Verklaring Omtrent Gedrag), "
    "niet ouder dan 12 maanden bij aanvang van de opdracht."
)
d.add_heading("Artikel 9 — Boeteclausules", level=2)
d.add_paragraph(
    "9.1 Bij het niet behalen van de in artikel 3 genoemde reactietijden wordt een boete van "
    "€250 per overschrijding in rekening gebracht, met een maximum van €5.000 per maand."
)

d.save(OUT)
print(f"SLA-testdocument opgeslagen: {OUT}")
