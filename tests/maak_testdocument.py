"""Genereert een klein synthetisch testdocument (offerte-aanvraag) voor end-to-end tests."""
from pathlib import Path

import docx

OUT = Path(__file__).parent / "sample_offerteaanvraag.docx"

d = docx.Document()
d.add_heading("Offerteaanvraag Schoonmaakdienstverlening – Kantoorlocatie Amsterdam", level=1)
d.add_paragraph(
    "Opdrachtgever vraagt inschrijvers een offerte uit te brengen voor dagelijkse "
    "schoonmaakdienstverlening op de hoofdlocatie te Amsterdam, ingaande 1 januari 2027, "
    "voor een periode van 3 jaar met optie tot verlenging van 2x1 jaar."
)
d.add_heading("Gunningscriteria (EMVI)", level=2)
d.add_paragraph(
    "De opdracht wordt gegund op basis van de Beste Prijs-Kwaliteitverhouding (BPKV). "
    "Kwaliteit weegt voor 60% en Prijs voor 40% mee in de eindscore."
)
d.add_heading("Kwaliteit (60%)", level=3)
d.add_paragraph(
    "1. Plan van aanpak: beschrijf hoe de dienstverlening wordt ingericht, inclusief "
    "kwaliteitscontrole en KPI-monitoring (weging 25%)."
)
d.add_paragraph(
    "2. Duurzaamheid: beschrijf het duurzaamheidsbeleid en gebruikte schoonmaakmiddelen "
    "(weging 15%)."
)
d.add_paragraph(
    "3. Ervaring: minimaal 3 referenties van vergelijkbare opdrachten in de afgelopen 5 jaar "
    "(weging 20%)."
)
d.add_heading("Prijs (40%)", level=3)
d.add_paragraph("De totale jaarprijs excl. BTW telt mee voor 40% van de eindscore.")
d.add_heading("Knock-outcriteria", level=3)
d.add_paragraph(
    "- Inschrijver dient te beschikken over een geldig ISO 9001-certificaat. Offertes zonder "
    "geldig certificaat worden uitgesloten van verdere beoordeling."
)
d.add_paragraph(
    "- Inschrijver dient een geldige VCA**-certificering te overleggen."
)

d.save(OUT)
print(f"Testdocument opgeslagen: {OUT}")
